"""
services/parser.py
─────────────────────────────────────────────────────────────────────────────
Calibr – Resume Parsing Service

Responsibilities:
  1. Extract plain text from uploaded PDF or DOCX resume files.
  2. Use Google Gemini (gemini-flash-latest) to intelligently extract skills
     from that raw text and return them as a clean Python list.
  3. Provide a single high-level parse_resume() function that the resume
     router can call with raw file bytes and a filename.

No embeddings happen here — that is handled by embedder.py.
─────────────────────────────────────────────────────────────────────────────
"""

# ── Standard library ──────────────────────────────────────────────────────────
import io
import json
import logging
import re

# ── Third-party: PDF parsing ──────────────────────────────────────────────────
from pypdf import PdfReader                           # reads PDF bytes page-by-page

# ── Third-party: DOCX parsing ─────────────────────────────────────────────────
from docx import Document                             # reads DOCX paragraphs

# ── Third-party: LangChain ────────────────────────────────────────────────────
from langchain_core.messages import HumanMessage

# ── Standard library: environment ─────────────────────────────────────────────
import os
from langchain_groq import ChatGroq

# ── Module-level logger ───────────────────────────────────────────────────────
logger = logging.getLogger("calibr.parser")


# ─────────────────────────────────────────────────────────────────────────────
#  Helper: build the Groq LLM client
# ─────────────────────────────────────────────────────────────────────────────
def get_llm():
    # Groq API - free tier, blazing fast
    # Get free key at console.groq.com
    return ChatGroq(
        model=os.getenv("GROQ_MODEL", 
                        "llama-3.1-8b-instant"),
        api_key=os.getenv("GROQ_API_KEY"),
        temperature=0.1
    )


# ─────────────────────────────────────────────────────────────────────────────
#  Function 1: extract_text_from_pdf
# ─────────────────────────────────────────────────────────────────────────────
def extract_text_from_pdf(file_bytes: bytes) -> str:
    """
    Extract all readable text from a PDF file supplied as raw bytes.

    Why bytes? FastAPI UploadFile gives us bytes via await file.read(),
    so we wrap them in io.BytesIO to make pypdf happy without writing to disk.

    Process:
        1. Wrap bytes in an in-memory BytesIO buffer.
        2. Open with PdfReader — handles multi-page PDFs automatically.
        3. Extract text from every page and join with newlines.
        4. Collapse multiple consecutive whitespace characters to a single space
           and strip leading/trailing whitespace.

    Args:
        file_bytes: Raw bytes of the PDF file.

    Returns:
        A single clean string containing all text across all pages.
        Returns an empty string if extraction fails (logged as error).
    """
    try:
        # Wrap raw bytes in a buffer so pypdf can seek through it
        buffer = io.BytesIO(file_bytes)

        # Open the PDF — PdfReader handles encrypted PDFs too (with no password)
        reader = PdfReader(buffer)

        page_texts = []
        for page_num, page in enumerate(reader.pages):
            # extract_text() can return None for image-only pages
            page_text = page.extract_text() or ""
            page_texts.append(page_text)
            logger.debug(f"PDF page {page_num + 1}: extracted {len(page_text)} chars")

        # Join all pages and normalise whitespace
        full_text = "\n".join(page_texts)
        full_text = re.sub(r"\s+", " ", full_text).strip()

        logger.info(f"PDF extraction complete — {len(full_text)} chars from {len(reader.pages)} pages")
        return full_text

    except Exception as e:
        logger.error(f"Failed to extract text from PDF: {e}")
        return ""  # return empty string so the caller can still proceed gracefully


# ─────────────────────────────────────────────────────────────────────────────
#  Function 2: extract_text_from_docx
# ─────────────────────────────────────────────────────────────────────────────
def extract_text_from_docx(file_bytes: bytes) -> str:
    """
    Extract all readable text from a DOCX file supplied as raw bytes.

    DOCX files store content in XML inside a ZIP archive. python-docx handles
    this transparently — we just wrap bytes in BytesIO and open with Document().

    Process:
        1. Wrap bytes in BytesIO.
        2. Open with python-docx Document.
        3. Iterate over all paragraphs, join their text with newlines.
        4. Strip extra whitespace.

    Args:
        file_bytes: Raw bytes of the DOCX file.

    Returns:
        A single clean string of the document's text content.
        Returns an empty string if extraction fails (logged as error).

    Note:
        Tables, headers, and footers in DOCX are not in paragraphs. If a
        resume uses heavy table formatting, some content may be missed.
        This is acceptable for the MVP.
    """
    try:
        buffer = io.BytesIO(file_bytes)

        # Open the document — python-docx unpacks the XML automatically
        doc = Document(buffer)

        # doc.paragraphs gives every paragraph element in order
        # paragraph.text is the plain string content of that paragraph
        para_texts = [para.text for para in doc.paragraphs if para.text.strip()]

        full_text = "\n".join(para_texts)
        full_text = re.sub(r"\s+", " ", full_text).strip()

        logger.info(f"DOCX extraction complete — {len(full_text)} chars from {len(doc.paragraphs)} paragraphs")
        return full_text

    except Exception as e:
        logger.error(f"Failed to extract text from DOCX: {e}")
        return ""


# ─────────────────────────────────────────────────────────────────────────────
#  Function 4: parse_resume  (main entry point for the resume router)
# ─────────────────────────────────────────────────────────────────────────────
def parse_resume(file_bytes: bytes, filename: str) -> dict:
    """
    High-level function that orchestrates the resume text extraction pipeline.

    Steps:
        1. Detect file type from the filename extension (.pdf or .docx).
        2. Call the appropriate extractor to get raw text.
        3. Return a dict with the extracted text.

    Args:
        file_bytes: Raw bytes of the uploaded resume file.
        filename  : Original filename from the upload.

    Returns:
        A dict with the key:
            {
                "raw_text": str  # full extracted text
            }

    Raises:
        ValueError: If the file extension is not .pdf or .docx.
    """
    name_lower = filename.lower().strip()

    logger.info(f"parse_resume called for file: {filename}")

    if name_lower.endswith(".pdf"):
        raw_text = extract_text_from_pdf(file_bytes)
    elif name_lower.endswith(".docx"):
        raw_text = extract_text_from_docx(file_bytes)
    else:
        raise ValueError(f"Unsupported file type: '{filename}'")

    if not raw_text:
        logger.warning(f"No text extracted from '{filename}'")
        return {"raw_text": ""}

    return {"raw_text": raw_text}
